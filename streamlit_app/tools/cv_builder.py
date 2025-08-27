import io
import re
from datetime import date
from typing import List, Dict, Any, Optional

import streamlit as st

# Optional deps (export)
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None

# Optional dep (import from PDF)
try:
    import PyPDF2
except Exception:
    PyPDF2 = None


# ─────────────────────────────────────────────────────────
# Presets
# ─────────────────────────────────────────────────────────
COUNTRY_PRESETS = {
    "International": {"date_fmt": "%Y-%m", "labels": {"phone": "Phone", "email": "Email", "location": "Location"}},
    "United States": {"date_fmt": "%b %Y", "labels": {"phone": "Phone", "email": "Email", "location": "Location"}},
    "United Kingdom": {"date_fmt": "%b %Y", "labels": {"phone": "Tel", "email": "Email", "location": "Location"}},
    "Netherlands / Belgium": {"date_fmt": "%m-%Y", "labels": {"phone": "Tel", "email": "E-mail", "location": "Locatie"}},
    "South Africa": {"date_fmt": "%b %Y", "labels": {"phone": "Cell", "email": "Email", "location": "Location"}},
}
DEFAULT_COUNTRY = "International"


# ─────────────────────────────────────────────────────────
# State init
# ─────────────────────────────────────────────────────────
def _init_state():
    ss = st.session_state
    ss.setdefault("cv_personal", {"full_name": "", "role": "", "email": "", "phone": "", "location": "",
                                  "website": "", "linkedin": "", "github": ""})
    ss.setdefault("cv_summary", "")
    ss.setdefault("cv_experience", [])
    ss.setdefault("cv_education", [])
    ss.setdefault("cv_skills", [])
    ss.setdefault("cv_links", [])
    ss.setdefault("cv_country", DEFAULT_COUNTRY)
    ss.setdefault("cv_template", "Clean")


# ─────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────
EMAIL_RE = re.compile(r'[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}', re.I)
PHONE_RE = re.compile(r'(\+?\d[\d\s\-\(\)]{7,}\d)')
URL_RE   = re.compile(r'(https?://\S+)', re.I)

def _format_date(d: date, fmt: str) -> str:
    try:
        return d.strftime(fmt)
    except Exception:
        return str(d)

def _docx_header(doc, text: str, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _lines(text: str) -> List[str]:
    if not text:
        return []
    if "," in text and "\n" not in text:
        parts = [p.strip() for p in text.split(",")]
    else:
        parts = [p.strip("• ").strip() for p in text.splitlines()]
    return [p for p in parts if p]


# ─────────────────────────────────────────────────────────
# Importers (DOCX / PDF / Markdown)
# ─────────────────────────────────────────────────────────
def _read_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    # include table text if any
    for tbl in doc.tables:
        for row in tbl.rows:
            texts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(t for t in texts if t and t.strip())

def _read_pdf(file_bytes: bytes) -> str:
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 not installed")
    bio = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(bio)
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            pass
    return "\n".join(out)

def _read_md(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""

SECTION_KEYS = ["summary", "skills", "experience", "education", "links"]

def _split_sections_from_text(text: str) -> Dict[str, str]:
    """
    Naive splitter: find headings like 'Summary', 'Skills', 'Experience', 'Education', 'Links'.
    Return text blocks for each section.
    """
    lines = [ln.rstrip() for ln in text.splitlines()]
    blocks: Dict[str, List[str]] = {k: [] for k in SECTION_KEYS}
    current = None

    for ln in lines:
        low = ln.strip().lower()
        if low in ("summary", "profile summary", "objective"):
            current = "summary";   continue
        if low.startswith("skills"):   current = "skills";   continue
        if low.startswith("experience") or low.startswith("work experience") or low.startswith("employment"):
            current = "experience";    continue
        if low.startswith("education"):
            current = "education";     continue
        if low.startswith("links") or low.startswith("projects") or low.startswith("portfolio"):
            current = "links";         continue
        if current:
            blocks[current].append(ln)

    return {k: "\n".join(v).strip() for k, v in blocks.items() if v}

def _extract_personal(text: str) -> Dict[str, str]:
    """
    Try to guess name (first non-empty line), email, phone, and links.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    email = EMAIL_RE.search(text)
    phone = PHONE_RE.search(text)
    urls  = URL_RE.findall(text)

    # Guess name: first non-empty line that isn't a heading keyword and not email/phone
    name = ""
    for ln in lines[:5]:  # only scan the top
        if EMAIL_RE.search(ln) or PHONE_RE.search(ln):
            continue
        low = ln.lower()
        if any(h in low for h in ["summary", "skills", "experience", "education", "links", "profile"]):
            continue
        # Avoid pure URL lines
        if URL_RE.search(ln):
            continue
        name = ln.strip()
        break

    # Guess role: next useful line after name
    role = ""
    if name:
        idx = lines.index(name) if name in lines else 0
        for ln in lines[idx+1: idx+4]:
            if EMAIL_RE.search(ln) or PHONE_RE.search(ln) or URL_RE.search(ln):
                continue
            low = ln.lower()
            if any(h in low for h in ["summary", "skills", "experience", "education", "links", "profile"]):
                break
            role = ln.strip()
            break

    website = ""
    linkedin = ""
    github = ""
    for u in urls:
        ul = u.lower()
        if "linkedin.com" in ul:
            linkedin = u
        elif "github.com" in ul:
            github = u
        elif not website:
            website = u

    return {
        "full_name": name,
        "role": role,
        "email": email.group(0) if email else "",
        "phone": phone.group(0) if phone else "",
        "location": "",
        "website": website,
        "linkedin": linkedin,
        "github": github,
    }

def _parse_skills(block: str) -> List[str]:
    if not block:
        return []
    # Prefer comma separated on first line; else bullets/lines
    first = block.splitlines()[0]
    if "," in first and len(block.splitlines()) <= 3:
        return _lines(first)
    return _lines(block)

def _parse_experience(block: str) -> List[Dict[str, Any]]:
    """
    Very simple parser: treats non-empty lines as bullets under a single role unless a line
    looks like a header (Company — Title) then starts a new item.
    You can refine later as needed.
    """
    if not block:
        return []
    exps: List[Dict[str, Any]] = []
    curr: Dict[str, Any] = {"title": "", "company": "", "start": date(2020,1,1), "end": date(2022,1,1), "current": False, "bullets": []}

    def push():
        nonlocal curr
        if curr["title"] or curr["company"] or curr["bullets"]:
            exps.append(curr)
        curr = {"title": "", "company": "", "start": date(2020,1,1), "end": date(2022,1,1), "current": False, "bullets": []}

    for ln in block.splitlines():
        t = ln.strip("• ").strip()
        if not t:
            continue
        # crude header detector: has a dash and capital letter word
        if ("—" in t or "-" in t) and any(w[:1].isupper() for w in t.split()):
            push()
            # split "Title — Company" or "Company — Title"
            parts = [p.strip() for p in re.split(r"[—\-]+", t, maxsplit=1)]
            if len(parts) == 2:
                # guess which is title: if contains words like 'Manager/Engineer/Designer' treat as title
                if re.search(r"(manager|engineer|designer|lead|developer|head|director|analyst)", parts[0], re.I):
                    curr["title"], curr["company"] = parts[0], parts[1]
                else:
                    curr["company"], curr["title"] = parts[0], parts[1]
            else:
                curr["title"] = t
        else:
            curr["bullets"].append(t)
    push()
    return exps

def _parse_education(block: str) -> List[Dict[str, Any]]:
    if not block:
        return []
    out: List[Dict[str, Any]] = []
    for ln in block.splitlines():
        t = ln.strip("• ").strip()
        if not t:
            continue
        # naive split "Degree — Institution"
        parts = [p.strip() for p in re.split(r"[—\-]+", t, maxsplit=1)]
        deg, inst = (parts + ["", ""])[:2]
        out.append({"degree": deg, "institution": inst, "start": date(2018,1,1), "end": date(2022,1,1), "notes": ""})
    return out

def _parse_links(block: str) -> List[Dict[str, str]]:
    if not block:
        return []
    out = []
    for ln in block.splitlines():
        ln = ln.strip()
        if not ln:
            continue
        m = URL_RE.search(ln)
        if m:
            url = m.group(0)
            label = ln.replace(url, "").strip(" :-") or "Link"
            out.append({"label": label, "url": url})
    return out

def import_cv_from_bytes(file_bytes: bytes, kind: str) -> Dict[str, Any]:
    """
    kind: 'docx' | 'pdf' | 'md'
    Returns a dict with keys: personal, summary, skills, experience, education, links
    """
    if kind == "docx":
        text = _read_docx(file_bytes)
    elif kind == "pdf":
        text = _read_pdf(file_bytes)
    elif kind == "md":
        text = _read_md(file_bytes)
    else:
        raise ValueError("Unsupported kind")

    personal = _extract_personal(text)
    blocks = _split_sections_from_text(text)

    summary = blocks.get("summary", "")
    skills = _parse_skills(blocks.get("skills", ""))
    experience = _parse_experience(blocks.get("experience", ""))
    education = _parse_education(blocks.get("education", ""))
    links = _parse_links(blocks.get("links", ""))

    return {
        "personal": personal,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": education,
        "links": links,
    }


# ─────────────────────────────────────────────────────────
# DOCX Export
# ─────────────────────────────────────────────────────────
def build_docx(cv: Dict[str, Any], template: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx not available")

    doc = Document()
    full_name = cv["personal"]["full_name"].strip() or "Your Name"
    role = cv["personal"]["role"].strip()

    # Header
    _docx_header(doc, full_name, size=22)
    if role:
        r = doc.add_paragraph().add_run(role)
        r.italic = True
        r.font.size = Pt(12)

    # Contact line
    labels = cv["labels"]
    contact_bits = []
    for k in ("email", "phone", "location", "website", "linkedin", "github"):
        val = cv["personal"].get(k) or ""
        if val:
            contact_bits.append(f"{labels.get(k, k.capitalize())}: {val}")
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    if template == "Clean":
        if cv["summary"]:
            _docx_header(doc, "Summary")
            doc.add_paragraph(cv["summary"])
        if cv["experience"]:
            _docx_header(doc, "Experience")
            for e in cv["experience"]:
                doc.add_paragraph(f"{e['title']} — {e['company']}")
                for b in e["bullets"]:
                    doc.add_paragraph(b, style="List Bullet")
        if cv["education"]:
            _docx_header(doc, "Education")
            for ed in cv["education"]:
                doc.add_paragraph(f"{ed['degree']} — {ed['institution']}")
        if cv["skills"]:
            _docx_header(doc, "Skills")
            doc.add_paragraph(", ".join(cv["skills"]))
        if cv["links"]:
            _docx_header(doc, "Links")
            for l in cv["links"]:
                doc.add_paragraph(f"{l['label']}: {l['url']}")

    elif template == "Compact":
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left, right = table.rows[0].cells

        if cv["skills"]:
            left.add_paragraph("Skills", "Heading 2")
            left.add_paragraph(", ".join(cv["skills"]))
        if cv["links"]:
            left.add_paragraph("Links", "Heading 2")
            for l in cv["links"]:
                left.add_paragraph(f"{l['label']}: {l['url']}")

        if cv["summary"]:
            right.add_paragraph("Summary", "Heading 2")
            right.add_paragraph(cv["summary"])
        if cv["experience"]:
            right.add_paragraph("Experience", "Heading 2")
            for e in cv["experience"]:
                right.add_paragraph(f"{e['title']} — {e['company']}")
                for b in e["bullets"]:
                    right.add_paragraph(b, style="List Bullet")
        if cv["education"]:
            right.add_paragraph("Education", "Heading 2")
            for ed in cv["education"]:
                right.add_paragraph(f"{ed['degree']} — {ed['institution']}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ─────────────────────────────────────────────────────────
# PDF Export (simple)
# ─────────────────────────────────────────────────────────
def build_pdf(cv: Dict[str, Any]) -> bytes:
    if canvas is None:
        raise RuntimeError("reportlab not available")

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    y = height - 50

    def line(txt, sz=10, gap=14, bold=False):
        nonlocal y
        c.setFont("Helvetica-Bold" if bold else "Helvetica", sz)
        c.drawString(40, y, txt)
        y -= gap

    full_name = cv["personal"]["full_name"] or "Your Name"
    line(full_name, 16, 20, True)
    if cv["personal"]["role"]:
        line(cv["personal"]["role"], 12, 16)

    if cv["summary"]:
        line("Summary", 12, 16, True)
        for p in cv["summary"].splitlines():
            line(p, 10, 14)

    c.showPage()
    c.save()
    return bio.getvalue()


# ─────────────────────────────────────────────────────────
# Build dict from session
# ─────────────────────────────────────────────────────────
def _build_cv_dict() -> Dict[str, Any]:
    preset = COUNTRY_PRESETS.get(st.session_state.cv_country, COUNTRY_PRESETS[DEFAULT_COUNTRY])
    return {
        "personal": st.session_state.cv_personal,
        "summary": st.session_state.cv_summary,
        "experience": st.session_state.cv_experience,
        "education": st.session_state.cv_education,
        "skills": st.session_state.cv_skills,
        "links": st.session_state.cv_links,
        "date_fmt": preset["date_fmt"],
        "labels": preset["labels"],
    }


# ─────────────────────────────────────────────────────────
# Streamlit UI
# ─────────────────────────────────────────────────────────
def render():
    st.header("📄 International CV Builder (with Import)")

    _init_state()
    st.sidebar.subheader("Settings")
    st.session_state.cv_country = st.sidebar.selectbox("Country / Locale", list(COUNTRY_PRESETS.keys()))
    st.session_state.cv_template = st.sidebar.radio("Template", ["Clean", "Compact"])

    # Upload & Import
    st.subheader("Upload & Import CV (optional)")
    up = st.file_uploader("Upload an existing CV to prefill (DOCX, PDF, Markdown)",
                          type=["docx", "pdf", "md", "markdown"], accept_multiple_files=False)
    if up:
        kind = "docx" if up.name.lower().endswith(".docx") else "pdf" if up.name.lower().endswith(".pdf") else "md"
        try:
            parsed = import_cv_from_bytes(up.read(), kind)
            with st.expander("Detected data (review before applying)", expanded=True):
                st.write("**Personal**", parsed["personal"])
                st.write("**Summary**", parsed["summary"])
                st.write("**Skills**", parsed["skills"])
                st.write("**Experience**", parsed["experience"])
                st.write("**Education**", parsed["education"])
                st.write("**Links**", parsed["links"])
            if st.button("✅ Apply to form"):
                # Merge non-destructively: only fill empty fields, append lists
                p = st.session_state.cv_personal
                for k, v in parsed["personal"].items():
                    if v and not p.get(k):
                        p[k] = v
                if parsed["summary"] and not st.session_state.cv_summary:
                    st.session_state.cv_summary = parsed["summary"]
                if parsed["skills"]:
                    st.session_state.cv_skills = list(dict.fromkeys(st.session_state.cv_skills + parsed["skills"]))
                if parsed["experience"]:
                    st.session_state.cv_experience.extend(parsed["experience"])
                if parsed["education"]:
                    st.session_state.cv_education.extend(parsed["education"])
                if parsed["links"]:
                    st.session_state.cv_links.extend(parsed["links"])
                st.success("Imported. Scroll down to review & export.")
        except Exception as e:
            st.error(f"Could not import file: {e}")

    # Personal info
    st.subheader("Personal information")
    p = st.session_state.cv_personal
    c1, c2, c3 = st.columns(3)
    with c1:
        p["full_name"] = st.text_input("Full name*", value=p["full_name"])
        p["email"] = st.text_input("Email", value=p["email"])
    with c2:
        p["role"] = st.text_input("Role / Headline", value=p["role"])
        p["phone"] = st.text_input(COUNTRY_PRESETS[st.session_state.cv_country]["labels"]["phone"], value=p["phone"])
    with c3:
        p["location"] = st.text_input(COUNTRY_PRESETS[st.session_state.cv_country]["labels"]["location"], value=p["location"])
        p["website"] = st.text_input("Website / Portfolio", value=p["website"])

    c4, c5 = st.columns(2)
    with c4:
        p["linkedin"] = st.text_input("LinkedIn", value=p["linkedin"])
    with c5:
        p["github"] = st.text_input("GitHub", value=p["github"])

    # Summary
    st.subheader("Summary")
    st.session_state.cv_summary = st.text_area("3–5 lines about yourself", value=st.session_state.cv_summary, height=120)

    # Experience / Education / Skills / Links (simple adders)
    st.divider()
    st.subheader("Skills")
    skills_add = st.text_input("Add a skill and press Enter", key="skill_add")
    if skills_add:
        st.session_state.cv_skills.append(skills_add.strip())
        st.session_state.skill_add = ""
    if st.session_state.cv_skills:
        cols = st.columns(6)
        for i, sk in enumerate(st.session_state.cv_skills):
            cols[i % 6].markdown(f"- {sk}")

    st.divider()
    st.subheader("Experience (quick bullets)")
    exp_add = st.text_area("Add bullets (one per line). For a new role header, write 'Title — Company' on its own line.",
                           key="exp_add")
    if st.button("Add to experience"):
        st.session_state.cv_experience.extend(_parse_experience(exp_add))
        st.session_state.exp_add = ""

    st.divider()
    st.subheader("Education (quick lines)")
    edu_add = st.text_area("Add lines (e.g., 'B.Sc. Computer Science — University of Cape Town')", key="edu_add")
    if st.button("Add to education"):
        st.session_state.cv_education.extend(_parse_education(edu_add))
        st.session_state.edu_add = ""

    st.divider()
    st.subheader("Links")
    link_label = st.text_input("Label", key="link_label")
    link_url = st.text_input("URL", key="link_url")
    if st.button("Add link"):
        if link_label and link_url:
            st.session_state.cv_links.append({"label": link_label, "url": link_url})
            st.session_state.link_label = ""
            st.session_state.link_url = ""

    # Preview & Export
    st.divider()
    st.subheader("Preview & Export")
    cv = _build_cv_dict()
    template = st.session_state.cv_template

    c1, c2 = st.columns(2)
    with c1:
        if Document:
            try:
                docx_bytes = build_docx(cv, template)
                st.download_button("⬇️ Download DOCX", docx_bytes, file_name=f"{cv['personal']['full_name'] or 'cv'}.docx")
            except Exception as e:
                st.error(f"DOCX error: {e}")
        else:
            st.info("DOCX export unavailable (python-docx not installed).")
    with c2:
        if canvas:
            try:
                pdf_bytes = build_pdf(cv)
                st.download_button("⬇️ Download PDF", pdf_bytes, file_name=f"{cv['personal']['full_name'] or 'cv'}.pdf")
            except Exception as e:
                st.error(f"PDF error: {e}")
        else:
            st.info("PDF export unavailable (reportlab not installed).")

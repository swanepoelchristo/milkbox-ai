import io
from datetime import date
from typing import List, Dict, Optional

import streamlit as st

# Optional deps for exports
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
except Exception:
    canvas = None


# ─────────────────────────────────────────────────────────
# Template definitions
# ─────────────────────────────────────────────────────────
TEMPLATES = {
    "Modern": {
        "docx": {
            "margins_in": (0.7, 0.7, 0.7, 0.7),  # (top, right, bottom, left)
            "name_size": 20,
            "role_size": 12,
            "h1_size": 13,
            "h2_size": 11,
            "body_size": 10,
            "line_space": 1.15,
            "divider": True,
        },
        "pdf": {
            "margins_mm": (18, 18, 20, 18),      # (top, right, bottom, left)
            "name_size": 18,
            "role_size": 11,
            "h_size": 11.5,
            "body_size": 9.8,
            "section_gap_mm": 6,
            "divider": True,
        },
    },
    "Classic": {
        "docx": {
            "margins_in": (1.0, 1.0, 1.0, 1.0),
            "name_size": 18,
            "role_size": 11,
            "h1_size": 12,
            "h2_size": 10.5,
            "body_size": 10,
            "line_space": 1.25,
            "divider": False,
        },
        "pdf": {
            "margins_mm": (22, 22, 22, 22),
            "name_size": 16,
            "role_size": 10.5,
            "h_size": 11,
            "body_size": 9.5,
            "section_gap_mm": 8,
            "divider": False,
        },
    },
}

LOCALES = [
    "International (DD MMM YYYY)",
    "US (Mon YYYY)",
    "UK (DD Mon YYYY)"
]


# ─────────────────────────────────────────────────────────
# State / Helpers
# ─────────────────────────────────────────────────────────
def _init_state() -> Dict:
    if "cvpro" not in st.session_state:
        st.session_state.cvpro = {
            "basics": {
                "full_name": "",
                "email": "",
                "phone": "",
                "location": "",
                "website": "",
                "linkedin": "",
                "summary": "",
                "locale": LOCALES[0],
                "template": "Modern",
            },
            "experience": [],  # {title, company, location, start, end, current, bullets}
            "education": [],   # {degree, school, location, start, end, bullets}
            "skills": [],      # [str]
        }
    return st.session_state.cvpro


def parse_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def parse_docx_text(file_bytes: bytes) -> str:
    if Document is None:
        return ""
    with io.BytesIO(file_bytes) as f:
        doc = Document(f)
    lines = [p.text for p in doc.paragraphs]
    return "\n".join([x for x in lines if x and x.strip()])


def prefill_from_text(text: str, data: Dict) -> None:
    # Very light heuristics
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return
    if not data["basics"]["full_name"] and len(lines[0].split()) <= 6:
        data["basics"]["full_name"] = lines[0]
    for ln in lines[:40]:
        ll = ln.lower()
        if "@" in ln and not data["basics"]["email"]:
            data["basics"]["email"] = ln
        if "linkedin" in ll and not data["basics"]["linkedin"]:
            data["basics"]["linkedin"] = ln
        if ("http" in ll or "www." in ll) and not data["basics"]["website"]:
            data["basics"]["website"] = ln


def fmt_date(d: Optional[date], locale: str) -> str:
    if not d:
        return ""
    if locale.startswith("US"):
        return d.strftime("%b %Y")
    if locale.startswith("UK"):
        return d.strftime("%d %b %Y")
    return d.strftime("%d %b %Y")


# ─────────────────────────────────────────────────────────
# DOCX builder with template styles
# ─────────────────────────────────────────────────────────
def _docx_set_margins(section, top_in, right_in, bottom_in, left_in):
    section.top_margin = Inches(top_in)
    section.right_margin = Inches(right_in)
    section.bottom_margin = Inches(bottom_in)
    section.left_margin = Inches(left_in)


def _docx_para(doc, text: str, size: int, bold=False, italic=False, space_after_pt=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after_pt)
    return p


def _docx_divider(doc):
    # add a thin horizontal rule
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(4)
    p_format.space_after = Pt(6)
    # Use bottom border on paragraph
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')       # thin
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom)
    p_pr.append(pbdr)


def make_docx(data: Dict) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx not installed (add python-docx)")

    doc = Document()
    b = data["basics"]
    t = TEMPLATES.get(b.get("template", "Modern"))["docx"]

    # Page margins
    _docx_set_margins(doc.sections[0], *t["margins_in"])

    # Name / role
    _docx_para(doc, b.get("full_name") or "Curriculum Vitae", t["name_size"], bold=True, space_after_pt=4)
    if b.get("role"):
        _docx_para(doc, b["role"], t["role_size"], italic=True, space_after_pt=6)

    # Contact line
    contact = " | ".join([x for x in [b.get("email"), b.get("phone"), b.get("location"), b.get("website"), b.get("linkedin")] if x])
    if contact:
        _docx_para(doc, contact, t["body_size"], space_after_pt=8)

    if t["divider"]:
        _docx_divider(doc)

    # Summary
    if b.get("summary"):
        _docx_para(doc, "Professional Summary", t["h1_size"], bold=True, space_after_pt=2)
        _docx_para(doc, b["summary"], t["body_size"], space_after_pt=8)
        if t["divider"]:
            _docx_divider(doc)

    # Experience
    if data["experience"]:
        _docx_para(doc, "Experience", t["h1_size"], bold=True, space_after_pt=2)
        for job in data["experience"]:
            title = job.get("title", "")
            company = job.get("company", "")
            head = f"{title} — {company}".strip(" —") or "Role"
            _docx_para(doc, head, t["h2_size"], bold=True, space_after_pt=0)

            start = fmt_date(job.get("start"), b["locale"])
            end = "Present" if job.get("current") else fmt_date(job.get("end"), b["locale"])
            meta = " | ".join([x for x in [job.get("location",""), f"{start} — {end}".strip(" —")] if x])
            if meta:
                _docx_para(doc, meta, t["body_size"], italic=True, space_after_pt=2)

            for bl in (job.get("bullets") or []):
                if bl.strip():
                    para = doc.add_paragraph(style=None)
                    run = para.add_run("• " + bl.strip())
                    run.font.size = Pt(t["body_size"])
                    para.paragraph_format.space_after = Pt(1)
        if t["divider"]:
            _docx_divider(doc)

    # Education
    if data["education"]:
        _docx_para(doc, "Education", t["h1_size"], bold=True, space_after_pt=2)
        for ed in data["education"]:
            head = f"{ed.get('degree','')} — {ed.get('school','')}".strip(" —") or "Education"
            _docx_para(doc, head, t["h2_size"], bold=True, space_after_pt=0)

            start = fmt_date(ed.get("start"), b["locale"])
            end = fmt_date(ed.get("end"), b["locale"])
            meta = " | ".join([x for x in [ed.get("location",""), f"{start} — {end}".strip(' —')] if x])
            if meta:
                _docx_para(doc, meta, t["body_size"], italic=True, space_after_pt=2)

            for bl in (ed.get("bullets") or []):
                if bl.strip():
                    para = doc.add_paragraph(style=None)
                    run = para.add_run("• " + bl.strip())
                    run.font.size = Pt(t["body_size"])
                    para.paragraph_format.space_after = Pt(1)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────
# PDF builder with template styles
# ─────────────────────────────────────────────────────────
def _pdf_divider(c, x, y, width_mm=170):
    c.line(x, y, x + width_mm * mm, y)

def wrap_text(s: str, width_chars: int) -> List[str]:
    words = s.split()
    out, cur = [], []
    for w in words:
        cur.append(w)
        if len(" ".join(cur)) > width_chars:
            out.append(" ".join(cur[:-1]))
            cur = [w]
    if cur:
        out.append(" ".join(cur))
    return out

def make_pdf(data: Dict) -> bytes:
    if canvas is None:
        raise RuntimeError("reportlab not installed (add reportlab)")

    b = data["basics"]
    tp = TEMPLATES.get(b.get("template", "Modern"))["pdf"]

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4

    top, right, bottom, left = tp["margins_mm"]
    x = left * mm
    y = H - top * mm
    body_w_chars = 95 if tp is TEMPLATES["Modern"]["pdf"] else 90

    def line(txt: str, size: float, dy_mm: float, bold=False):
        nonlocal y
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(x, y, txt)
        y -= dy_mm * mm

    # Name / role
    line(b.get("full_name") or "Curriculum Vitae", tp["name_size"], 7, bold=True)
    if b.get("role"):
        line(b["role"], tp["role_size"], 6)

    # Contact
    contact = " | ".join([x for x in [b.get("email"), b.get("phone"), b.get("location"), b.get("website"), b.get("linkedin")] if x])
    if contact:
        line(contact, tp["body_size"], 6)

    if tp["divider"]:
        _pdf_divider(c, x, y + 2, width_mm=(W/mm - left - right))
        y -= 2 * mm

    # Summary
    if b.get("summary"):
        line("Professional Summary", tp["h_size"], tp["section_gap_mm"], bold=True)
        for chunk in wrap_text(b["summary"], body_w_chars):
            line(chunk, tp["body_size"], 5)

        if tp["divider"]:
            _pdf_divider(c, x, y + 2, width_mm=(W/mm - left - right))
            y -= 2 * mm

    # Experience
    if data["experience"]:
        line("Experience", tp["h_size"], tp["section_gap_mm"], bold=True)
        for job in data["experience"]:
            title = job.get("title",""); company = job.get("company","")
            start = fmt_date(job.get("start"), b["locale"])
            end = "Present" if job.get("current") else fmt_date(job.get("end"), b["locale"])
            head = f"{title} — {company}".strip(" —") or "Role"
            line(head, tp["h_size"], 5.5, bold=True)

            meta = " | ".join([x for x in [job.get("location",""), f"{start} — {end}".strip(" —")] if x])
            if meta:
                line(meta, tp["body_size"], 5)

            for bl in (job.get("bullets") or []):
                for chunk in wrap_text("• " + bl.strip(), body_w_chars):
                    line(chunk, tp["body_size"], 5)

        if tp["divider"]:
            _pdf_divider(c, x, y + 2, width_mm=(W/mm - left - right))
            y -= 2 * mm

    # Education
    if data["education"]:
        line("Education", tp["h_size"], tp["section_gap_mm"], bold=True)
        for ed in data["education"]:
            head = f"{ed.get('degree','')} — {ed.get('school','')}".strip(" —") or "Education"
            line(head, tp["h_size"], 5.5, bold=True)

            start = fmt_date(ed.get("start"), b["locale"])
            end = fmt_date(ed.get("end"), b["locale"])
            meta = " | ".join([x for x in [ed.get("location",""), f"{start} — {end}".strip(" —")] if x])
            if meta:
                line(meta, tp["body_size"], 5)

            for bl in (ed.get("bullets") or []):
                for chunk in wrap_text("• " + bl.strip(), body_w_chars):
                    line(chunk, tp["body_size"], 5)

    c.showPage()
    c.save()
    return buf.getvalue()


# ─────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────
def render():
    st.header("International CV Builder (Pro)")
    st.caption("Choose a template, fill your info, and export as DOCX/PDF. You can import a TXT/DOCX to prefill basics.")

    data = _init_state()
    b = data["basics"]

    with st.expander("📥 Import existing CV (TXT/DOCX)", expanded=False):
        up = st.file_uploader("Upload a TXT or DOCX file", type=["txt", "docx"])
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("Pre-fill"):
                if up:
                    raw = up.read()
                    txt = parse_txt(raw) if up.name.lower().endswith(".txt") else parse_docx_text(raw)
                    if txt:
                        prefill_from_text(txt, data)
                        st.success("Prefilled basic details.")
                    else:
                        st.warning("Could not read text from file.")
                else:
                    st.warning("Upload a TXT or DOCX file first.")
        with col2:
            st.write("This lightly extracts name/contact lines. You can edit everything below.")

    # Tabs
    t1, t2, t3, t4, t5 = st.tabs(["Basics & Template", "Experience", "Education", "Skills & Summary", "Export"])

    with t1:
        b["full_name"] = st.text_input("Full name", b["full_name"])
        b["role"] = st.text_input("Role / Headline", b.get("role", ""))
        c1, c2, c3 = st.columns(3)
        with c1:
            b["email"] = st.text_input("Email", b["email"])
            b["location"] = st.text_input("Location", b["location"])
        with c2:
            b["phone"] = st.text_input("Phone", b["phone"])
            b["website"] = st.text_input("Website/Portfolio", b["website"])
        with c3:
            b["linkedin"] = st.text_input("LinkedIn URL", b["linkedin"])
            b["locale"] = st.selectbox("Locale / Date format", LOCALES, index=LOCALES.index(b["locale"]) if b.get("locale") else 0)
        st.markdown("**Template**")
        b["template"] = st.radio("Choose template", list(TEMPLATES.keys()), horizontal=True, index=list(TEMPLATES.keys()).index(b.get("template","Modern")))

    with t2:
        st.subheader("Experience")
        if not data["experience"]:
            data["experience"].append({
                "title": "", "company": "", "location": "",
                "start": None, "end": None, "current": False, "bullets": []
            })
        for i, job in enumerate(data["experience"]):
            st.markdown(f"##### Role {i+1}")
            c1, c2 = st.columns(2)
            with c1:
                job["title"] = st.text_input("Job title", job["title"], key=f"xtitle{i}")
                job["company"] = st.text_input("Company", job["company"], key=f"xcomp{i}")
                job["location"] = st.text_input("Location", job["location"], key=f"xloc{i}")
            with c2:
                job["start"] = st.date_input("Start date", value=job["start"] or date(2021,1,1), key=f"xstart{i}")
                job["current"] = st.checkbox("I currently work here", value=job.get("current", False), key=f"xcur{i}")
                if not job["current"]:
                    job["end"] = st.date_input("End date", value=job["end"] or date.today(), key=f"xend{i}")
                else:
                    job["end"] = None
            new_b = st.text_area("Add bullet (one per add)", key=f"xbul{i}")
            cc1, cc2 = st.columns(2)
            if cc1.button("➕ Add bullet", key=f"xadd{i}") and new_b.strip():
                (job["bullets"] or []).append(new_b.strip())
                st.experimental_rerun()
            if job.get("bullets"):
                st.write("**Bullets:**")
                for bi, bl in enumerate(job["bullets"]):
                    colL, colR = st.columns([8,1])
                    with colL: st.write("- " + bl)
                    with colR:
                        if st.button("🗑", key=f"xdel{i}-{bi}"):
                            job["bullets"].pop(bi)
                            st.experimental_rerun()
            st.divider()
        if st.button("➕ Add another role"):
            data["experience"].append({
                "title": "", "company": "", "location": "",
                "start": None, "end": None, "current": False, "bullets": []
            })

    with t3:
        st.subheader("Education")
        if not data["education"]:
            data["education"].append({
                "degree": "", "school": "", "location": "",
                "start": None, "end": None, "bullets": []
            })
        for i, ed in enumerate(data["education"]):
            st.markdown(f"##### Program {i+1}")
            c1, c2 = st.columns(2)
            with c1:
                ed["degree"] = st.text_input("Degree / Program", ed["degree"], key=f"edeg{i}")
                ed["school"] = st.text_input("School", ed["school"], key=f"esch{i}")
                ed["location"] = st.text_input("Location", ed["location"], key=f"eloc{i}")
            with c2:
                ed["start"] = st.date_input("Start date", value=ed["start"] or date(2018,1,1), key=f"estart{i}")
                ed["end"] = st.date_input("End date", value=ed["end"] or date.today(), key=f"eend{i}")
            new_b = st.text_area("Add bullet (one per add)", key=f"ebul{i}")
            cc1, cc2 = st.columns(2)
            if cc1.button("➕ Add bullet", key=f"eadd{i}") and new_b.strip():
                (ed["bullets"] or []).append(new_b.strip())
                st.experimental_rerun()
            if ed.get("bullets"):
                st.write("**Bullets:**")
                for bi, bl in enumerate(ed["bullets"]):
                    colL, colR = st.columns([8,1])
                    with colL: st.write("- " + bl)
                    with colR:
                        if st.button("🗑", key=f"edel{i}-{bi}"):
                            ed["bullets"].pop(bi)
                            st.experimental_rerun()
            st.divider()
        if st.button("➕ Add another education"):
            data["education"].append({
                "degree": "", "school": "", "location": "",
                "start": None, "end": None, "bullets": []
            })

    with t4:
        st.subheader("Skills & Summary")
        new_skill = st.text_input("Add a skill")
        if st.button("➕ Add skill") and new_skill.strip():
            data["skills"].append(new_skill.strip())
        if data["skills"]:
            st.write("**Skills:** " + ", ".join(data["skills"]))
            if st.button("Clear skills"):
                data["skills"] = []
        b["summary"] = st.text_area("Professional summary", b["summary"], height=140)

    with t5:
        st.subheader("Export")
        colA, colB = st.columns(2)
        with colA:
            if st.button("⤓ Generate DOCX"):
                try:
                    docx_bytes = make_docx(data)
                    st.download_button("Download CV (DOCX)", data=docx_bytes, file_name="CV_Pro.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(str(e))
        with colB:
            if st.button("⤓ Generate PDF"):
                try:
                    pdf_bytes = make_pdf(data)
                    st.download_button("Download CV (PDF)", data=pdf_bytes, file_name="CV_Pro.pdf", mime="application/pdf")
                except Exception as e:
                    st.error(str(e))

    st.success("Template applied. Your choices persist while editing.")
